#!/usr/bin/env python3
"""
Markdown to DOCX Converter
Converts markdown documents to Microsoft Word format with Korean font support,
inline formatting (bold/italic), and nested list indentation.
"""

from docx import Document
from docx.shared import Pt, Inches
import re
import sys


class DOCXGenerator:
    """Generate DOCX files from markdown with Korean font support."""

    def __init__(self):
        """Initialize document with Korean-friendly font settings."""
        self.doc = Document()
        self.setup_korean_fonts()

    def setup_korean_fonts(self):
        """Set default Korean-friendly fonts (Malgun Gothic)."""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Malgun Gothic'
        font.size = Pt(11)

    def process_inline_formatting(self, text, paragraph):
        """
        Process inline formatting like bold and italic.

        Handles:
        - **bold** or __bold__
        - *italic* or _italic_

        Args:
            text: Text with markdown formatting
            paragraph: DOCX paragraph object to add runs to
        """
        # Split by ** for bold (handles both ** and __)
        parts = re.split(r'(\*\*.*?\*\*|__.*?__)', text)
        for part in parts:
            if (part.startswith('**') and part.endswith('**')) or \
               (part.startswith('__') and part.endswith('__')):
                # Bold text
                clean_text = part[2:-2]
                run = paragraph.add_run(clean_text)
                run.bold = True
                run.font.name = 'Malgun Gothic'
            elif part:
                # Regular text (may contain italic)
                # TODO: Add italic support if needed
                run = paragraph.add_run(part)
                run.font.name = 'Malgun Gothic'

    def get_indent_level(self, line):
        """
        Calculate indent level based on leading spaces.

        Args:
            line: Text line with possible leading spaces

        Returns:
            int: Indent level (2 spaces = 1 level)
        """
        spaces = len(line) - len(line.lstrip(' '))
        return spaces // 2

    def add_list_item(self, line, style='List Bullet'):
        """
        Add a list item with proper indentation.

        Args:
            line: List item line (with markers and spaces)
            style: DOCX list style ('List Bullet' or 'List Number')

        Returns:
            Paragraph object
        """
        indent_level = self.get_indent_level(line)
        text = line.lstrip()

        # Remove list markers
        if text.startswith('- ') or text.startswith('* '):
            text = text[2:]
        elif re.match(r'^\d+\. ', text):
            text = re.sub(r'^\d+\. ', '', text)

        # Create paragraph with style
        p = self.doc.add_paragraph(style=style)
        self.process_inline_formatting(text, p)

        # Apply indentation
        if indent_level > 0:
            p.paragraph_format.left_indent = Inches(0.5 * indent_level)

        return p

    def convert_markdown(self, content):
        """
        Convert markdown content to DOCX document.

        Args:
            content: Markdown text content

        Returns:
            Document object
        """
        # Skip frontmatter if present
        if content.startswith('---'):
            parts = content.split('---', 2)
            if len(parts) >= 3:
                content = parts[2].strip()

        lines = content.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i].rstrip()

            # Skip empty lines
            if not line:
                i += 1
                continue

            # Headers (must check from most specific to least)
            if line.lstrip().startswith('#### '):
                heading_text = line.lstrip()[5:]
                p = self.doc.add_heading(level=4)
                self.process_inline_formatting(heading_text, p)

            elif line.lstrip().startswith('### '):
                heading_text = line.lstrip()[4:]
                p = self.doc.add_heading(level=3)
                self.process_inline_formatting(heading_text, p)

            elif line.lstrip().startswith('## '):
                heading_text = line.lstrip()[3:]
                p = self.doc.add_heading(level=2)
                self.process_inline_formatting(heading_text, p)

            elif line.lstrip().startswith('# '):
                heading_text = line.lstrip()[2:]
                p = self.doc.add_heading(level=1)
                self.process_inline_formatting(heading_text, p)

            # Checkboxes
            elif line.lstrip().startswith('- [ ] '):
                indent_level = self.get_indent_level(line)
                p = self.doc.add_paragraph(style='List Bullet')
                p.add_run('☐ ')
                self.process_inline_formatting(line.lstrip()[6:], p)
                if indent_level > 0:
                    p.paragraph_format.left_indent = Inches(0.5 * indent_level)

            elif line.lstrip().startswith('- [x] ') or line.lstrip().startswith('- [X] '):
                indent_level = self.get_indent_level(line)
                p = self.doc.add_paragraph(style='List Bullet')
                p.add_run('☑ ')
                self.process_inline_formatting(line.lstrip()[6:], p)
                if indent_level > 0:
                    p.paragraph_format.left_indent = Inches(0.5 * indent_level)

            # Bullet lists
            elif line.lstrip().startswith('- ') or line.lstrip().startswith('* '):
                self.add_list_item(line, style='List Bullet')

            # Numbered lists
            elif re.match(r'^\s*\d+\. ', line):
                self.add_list_item(line, style='List Number')

            # Block quotes
            elif line.startswith('> '):
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                run = p.add_run(line[2:])
                run.italic = True
                run.font.name = 'Malgun Gothic'

            # Comments (skip)
            elif line.startswith('%%'):
                # Skip until closing %%
                while i < len(lines) and not lines[i].rstrip().endswith('%%'):
                    i += 1
                i += 1
                continue

            # Regular paragraphs
            else:
                if line.strip():
                    p = self.doc.add_paragraph()
                    self.process_inline_formatting(line, p)

            i += 1

        return self.doc

    def save(self, output_path):
        """
        Save document to file.

        Args:
            output_path: Path to save DOCX file
        """
        self.doc.save(output_path)


def create_docx_from_markdown(markdown_content, output_path, title=None):
    """
    Convert markdown content to DOCX file.

    Main entry point for the skill.

    Args:
        markdown_content: Markdown text content
        output_path: Path to save DOCX file
        title: Optional document title (not currently used)

    Returns:
        str: Path to created DOCX file
    """
    generator = DOCXGenerator()
    generator.convert_markdown(markdown_content)
    generator.save(output_path)
    return output_path


def main():
    """Command-line interface for testing."""
    if len(sys.argv) < 3:
        print("Usage: python docx_generator.py <input.md> <output.docx>")
        sys.exit(1)

    input_path = sys.argv[1]
    output_path = sys.argv[2]

    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()

    result = create_docx_from_markdown(content, output_path)
    print(f"✅ Successfully created: {result}")


if __name__ == '__main__':
    main()
